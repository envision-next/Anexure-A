"""
Fills the Annexure A Word template for one person and returns the .docx bytes.

Called by app.py:  build_document(base_dir, person, chosen_orgs, today_str)

The original template is never modified -- it is opened, filled in memory, and
saved to a BytesIO buffer.

Template layout (verified against the sample file):
  paragraphs:
    "Name: ..."                 -> person name
    "DIN/DPIN (If applicable):" -> person din
    after "(Signature)"         -> signature name (person name again)
    "Date: ..."                 -> today's date
    "Place: ..."                -> removed
  tables[0] Organizations : S.No | Name | Address | RERA Number
  tables[1] Project status: S.No | RERA Number | Completion | Complaints | Warrants | Revoked
"""

import io
import os
import glob
import copy

from docx import Document


def _find_template(base_dir):
    """The Annexure .docx in the folder (skip Word temp files ~$...)."""
    docs = [p for p in glob.glob(os.path.join(base_dir, "*.docx"))
            if not os.path.basename(p).startswith("~$")]
    if not docs:
        raise FileNotFoundError("No .docx template found in the folder.")
    # Prefer one that mentions Annexure; otherwise the first.
    for p in docs:
        if "annexure" in os.path.basename(p).lower():
            return p
    return docs[0]


def _set_para_text(paragraph, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _set_value_after_label(paragraph, label, value):
    """For 'Label: old' paragraphs -> 'Label: value'."""
    _set_para_text(paragraph, f"{label} {value}".rstrip())


def _delete_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def _set_cell(cell, text):
    """Write text into a cell, keeping the first run's formatting."""
    para = cell.paragraphs[0]
    runs = para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(text)
    # Clear any extra paragraphs in the cell.
    for extra in cell.paragraphs[1:]:
        _delete_paragraph(extra)


def _resize_table(table, header_rows, needed):
    """Ensure the table has exactly `needed` data rows, cloning the first data
    row so formatting/borders are preserved. Returns the list of data rows."""
    data_rows = table.rows[header_rows:]
    if not data_rows:
        raise ValueError("Template table has no data row to clone.")

    template_tr = data_rows[0]._tr

    # Add rows if we need more.
    while len(table.rows) - header_rows < needed:
        new_tr = copy.deepcopy(template_tr)
        table._tbl.append(new_tr)

    # Remove rows if we have too many.
    while len(table.rows) - header_rows > needed:
        last = table.rows[-1]._tr
        last.getparent().remove(last)

    return table.rows[header_rows:]


def _swap_yes_for_no(paragraph):
    """In the 'Are you a Director/...? - Yes' paragraph, turn 'Yes' into 'No'."""
    for run in paragraph.runs:
        if run.text.strip().lower() == "yes":
            run.text = run.text.replace("Yes", "No").replace("yes", "No")


def build_document(base_dir, person, orgs, today_str):
    template_path = _find_template(base_dir)
    doc = Document(template_path)

    # When a person belongs to no organization, the answer becomes "No" and
    # both tables get a single row of "NA" across every column.
    na = not orgs
    org_source = orgs if orgs else [None]

    # --- Header paragraphs ---
    seen_signature = False
    for p in doc.paragraphs:
        t = p.text.strip()
        low = t.lower()
        if low.startswith("name:"):
            _set_value_after_label(p, "Name:", person["name"])
        elif low.startswith("din/dpin"):
            _set_value_after_label(p, "DIN/DPIN (If applicable):", person["din"])
        elif low.startswith("date:"):
            _set_value_after_label(p, "Date:", today_str)
        elif low.startswith("place:"):
            _delete_paragraph(p)
        elif na and low.startswith("are you a director"):
            _swap_yes_for_no(p)
        elif "(signature)" in low:
            seen_signature = True
        elif seen_signature and t:
            # First non-empty paragraph after "(Signature)" = signature name.
            _set_para_text(p, person["name"])
            seen_signature = False

    # --- Table 0: Organizations ---
    org_table = doc.tables[0]
    rows = _resize_table(org_table, header_rows=1, needed=len(org_source))
    for i, (row, org) in enumerate(zip(rows, org_source), start=1):
        cells = row.cells
        if na:
            vals = ["NA", "NA", "NA", "NA"]
        else:
            vals = [str(i), org["name"], org["address"], org["rera"]]
        for cell, val in zip(cells, vals):
            _set_cell(cell, val)

    # --- Table 1: Project status ---
    status_table = doc.tables[1]
    rows = _resize_table(status_table, header_rows=1, needed=len(org_source))
    for i, (row, org) in enumerate(zip(rows, org_source), start=1):
        cells = row.cells
        if na:
            vals = ["NA", "NA", "NA", "NA", "NA", "NA"]
        else:
            vals = [str(i), org["rera"], org["completion"],
                    org["complaints"], org["warrants"], org["revoked"]]
        for cell, val in zip(cells, vals):
            _set_cell(cell, val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
